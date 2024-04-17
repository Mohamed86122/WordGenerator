from docx import Document
from docx.shared import Pt

# Create a new Word document
doc = Document()
doc.add_heading('Rapport d\'Installation et Configuration d\'Apache Hadoop', level=1)

# Introduction
doc.add_heading('1. Introduction', level=2)
doc.add_paragraph(
    "Ce rapport détaille les étapes nécessaires pour télécharger, installer, et configurer Apache Hadoop, "
    "un framework logiciel pour le stockage et le traitement distribué de grands ensembles de données. "
    "Le but est de fournir une documentation claire et concise pour faciliter la mise en place d'un environnement Hadoop."
)

# Téléchargement et Installation
doc.add_heading('2. Téléchargement et Installation de Hadoop', level=2)
doc.add_heading('Question', level=3)
doc.add_paragraph("Comment télécharger et installer Apache Hadoop ?")
doc.add_heading('Réponse et Explication', level=3)
doc.add_paragraph(
    "Apache Hadoop peut être téléchargé depuis le site officiel d'Apache. Il faut choisir la version appropriée, "
    "télécharger le fichier compressé, et le décompresser dans un répertoire spécifique sans espaces dans les noms de chemin. "
    "Cela prépare le système pour les étapes de configuration suivantes."
)
doc.add_heading('Commande', level=3)
doc.add_paragraph("Sur Linux/macOS, pour décompresser le fichier :")
doc.add_paragraph("tar -xzvf hadoop-x.y.z-src.tar.gz -C /chemin/vers/destination", style='Intense Quote')

# Configuration de Hadoop
doc.add_heading('3. Configuration de Hadoop', level=2)
doc.add_heading('Question', level=3)
doc.add_paragraph("Comment configurer les variables d'environnement et les fichiers pour Hadoop ?")
doc.add_heading('Réponse et Explication', level=3)
doc.add_paragraph(
    "La configuration de Hadoop nécessite de définir la variable d'environnement HADOOP_HOME et de mettre à jour le PATH. "
    "Il est également nécessaire de modifier les fichiers de configuration tels que core-site.xml, hdfs-site.xml, "
    "mapred-site.xml, et yarn-site.xml pour définir correctement les chemins et les paramètres du cluster."
)
doc.add_heading('Commande', level=3)
doc.add_paragraph("Exemple de commandes pour définir les variables d'environnement sur Linux/macOS :")
doc.add_paragraph("export HADOOP_HOME=/chemin/vers/hadoop\nexport PATH=$PATH:$HADOOP_HOME/bin:$HADOOP_HOME/sbin", style='Intense Quote')

# Formatage du Namenode
doc.add_heading('4. Formatage du Namenode', level=2)
doc.add_heading('Question', level=3)
doc.add_paragraph("Pourquoi et comment formater le Namenode ?")
doc.add_heading('Réponse et Explication', level=3)
doc.add_paragraph(
    "Le formatage du Namenode est crucial pour initialiser le système de fichiers HDFS, le préparant pour stocker des données. "
    "Cela doit être fait avant de démarrer Hadoop pour la première fois ou si une réinitialisation complète est nécessaire."
)
doc.add_heading('Commande', level=3)
doc.add_paragraph("Commande pour formater le Namenode :")
doc.add_paragraph("hadoop namenode -format", style='Intense Quote')

# Démarrage de Hadoop
doc.add_heading('5. Démarrage de Hadoop', level=2)
doc.add_heading('Question', level=3)
doc.add_paragraph("Comment démarrer les services Hadoop ?")
doc.add_heading('Réponse et Explication', level=3)
doc.add_paragraph(
    "Après la configuration et le formatage, les services Hadoop tels que Namenode, Datanode, ResourceManager, et NodeManager "
    "doivent être démarrés pour que le cluster soit opérationnel."
)
doc.add_heading('Commande', level=3)
doc.add_paragraph("Commandes pour démarrer Hadoop sur Linux/macOS :")
doc.add_paragraph("./start-all.sh", style='Intense Quote')

# Vérification des Services
doc.add_heading('6. Vérification des Services', level=2)
doc.add_heading('Question', level=3)
doc.add_paragraph("Comment vérifier que les services Hadoop fonctionnent correctement ?")
doc.add_heading('Réponse et Explication', level=3)
doc.add_paragraph(
    "Les interfaces web du Namenode et du ResourceManager permettent de vérifier le bon fonctionnement des services. "
    "Elles fournissent des informations sur la gestion des tâches et des ressources, ainsi que l'état général du cluster. L'accès à ces interfaces se fait via des URLs spécifiques."
)
doc.add_heading('URLs pour la Vérification', level=3)
doc.add_paragraph("Namenode : http://localhost:9870\nResourceManager : http://localhost:8088", style='Intense Quote')

# Conclusion
doc.add_heading('7. Conclusion', level=2)
doc.add_paragraph(
    "Ce rapport a guidé à travers les différentes étapes pour l'installation et la configuration de Hadoop. Chaque étape a été "
    "détaillée avec des explications et des commandes spécifiques pour assurer une compréhension claire et faciliter la mise en "
    "œuvre de Hadoop dans un environnement de développement ou de production. Le suivi de ces instructions devrait permettre "
    "de configurer un cluster Hadoop fonctionnel et prêt à l'emploi."
)
file_path = "C:\\Users\\Lenovo\\Desktop\\TP1 Hadoop\\Rapport.docx"
doc.save(file_path)
file_path
